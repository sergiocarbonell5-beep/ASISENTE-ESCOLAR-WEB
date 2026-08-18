# -*- coding: utf-8 -*-
"""
SISTEMA INTEGRAL EDUCATIVO — C.E.R. SIRAVITA
=======================================================================================
- Carga Ultra Rápida Dinámica (Sin bloqueos de pantalla).
- Evaluación Continua con Nombre y Fecha por Actividad (C1 - C10).
- Generación de Boletines Oficiales en PDF (2 Páginas con Escudo y Bandera).
- Adaptador Automático de Guías de Aprendizaje (PDF / Word / HTML a Escuela Nueva).
- Control de Asistencia en PDF y Excel.
- Dashboard Estadístico, Observador de Convivencia y Tabla de Líderes.
"""

import os
import re
import io
import base64
import datetime
import random
import urllib.parse
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

# Excel
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

# PDF y Word
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
import docx
from pypdf import PdfReader

from supabase import create_client, Client

# ================================================================
# CONFIGURACIÓN DE PÁGINA Y ESTILOS
# ================================================================
st.set_page_config(
    page_title="Asistente Educativo - C.E.R. Siravita",
    page_icon="🏫",
    layout="wide",
    initial_sidebar_state="expanded"
)

APP_TITLE = "Asistente Educativo Sergio Carbonell"
NOMBRE_ESCUELA = "C.E.R. Siravita"
SEDE_DEFECTO = "Chicago Alto"

PUNTOS_BASE = 10
PUNTOS_EXTRA_PUNTUALIDAD = 5
CUPOS_PUNTUALIDAD = 5

MENSAJES_ANIMO = [
    "¡Sigue así, campeón!",
    "¡Cada día sumas más!",
    "¡Eres un ejemplo para tus compañeros!",
    "¡La constancia es la clave del éxito!",
    "¡Un día más, un paso más cerca de tu meta!",
    "¡Tu esfuerzo no pasa desapercibido!",
]

MATERIAS_LISTA = [
    "Ciencias Naturales Educación ambiental",
    "Ciencias Sociales, historia, geografía, constitución y democrática.",
    "Cátedra de la Paz",
    "Educación Artística",
    "Educación Física, recreación y deportes",
    "Matemáticas",
    "Humanidades lengua castellana",
    "Idioma extranjero",
    "Tecnología e informática",
    "Ética y valores humanos",
    "Educación religiosa"
]

DIMENSIONES_EVALUACION = [
    "🗣️ Participación y Preguntas en Clase",
    "📚 Tareas y Guías de Aprendizaje",
    "📝 Evaluaciones y Pruebas Escritas",
    "🤝 Comportamiento y Actitud"
]

MESES_ESPANOL = [
    "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", 
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"
]

# Estilo para ajustar color institucional rápido de Streamlit
st.markdown("""
    <style>
    .stSpinner > div {
        border-top-color: #008037 !important;
    }
    </style>
""", unsafe_allow_html=True)

# ================================================================
# CONEXIÓN CON SUPABASE
# ================================================================
@st.cache_resource
def init_supabase() -> Client:
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    return create_client(url, key)

supabase = init_supabase()

if "profesor" not in st.session_state:
    st.session_state.profesor = None

# ================================================================
# LÓGICA DE ESCALA DE VALORACIÓN INSTITUCIONAL
# ================================================================
def obtener_valoracion_cualitativa(nota):
    if nota >= 4.8:
        return "DESEMPEÑO SUPERIOR", "EXCELENTE"
    elif nota >= 4.0:
        return "DESEMPEÑO ALTO", "BUENO"
    elif nota >= 3.0:
        return "DESEMPEÑO BÁSICO", "ACEPTABLE"
    elif nota > 0.0:
        return "DESEMPEÑO BAJO", "INSUFICIENTE"
    else:
        return "PENDIENTE", "PENDIENTE"

def normalizar_texto(txt):
    if not txt: return ""
    return re.sub(r'[^a-zA-Z0-9]', '', txt.lower())

# ================================================================
# AUTENTICACIÓN
# ================================================================
def registrar_profesor(nombre, email, password):
    email_clean = email.strip().lower()
    res = supabase.table("profesores").select("id").eq("email", email_clean).execute()
    if res.data:
        return False, "El correo electrónico ya está registrado."
    
    data = {"nombre": nombre.strip(), "email": email_clean, "password": password.strip()}
    supabase.table("profesores").insert(data).execute()
    return True, "¡Profesor registrado exitosamente! Ya puedes iniciar sesión."

def login_profesor(email, password):
    email_clean = email.strip().lower()
    res = supabase.table("profesores").select("*").eq("email", email_clean).eq("password", password.strip()).execute()
    if res.data:
        return res.data[0]
    return None

# ================================================================
# CONSULTAS A LA NUBE
# ================================================================
def obtener_grados(profesor_id):
    res = supabase.table("grados").select("*").eq("profesor_id", profesor_id).order("nombre").execute()
    return res.data

def agregar_grado(nombre, profesor_id):
    supabase.table("grados").insert({"nombre": nombre, "profesor_id": profesor_id}).execute()

def editar_grado(grado_id, nuevo_nombre):
    supabase.table("grados").update({"nombre": nuevo_nombre}).eq("id", grado_id).execute()

def eliminar_grado(grado_id):
    supabase.table("grados").delete().eq("id", grado_id).execute()

def obtener_estudiantes(grado_id, profesor_id):
    res = supabase.table("estudiantes").select("*").eq("grado_id", grado_id).eq("profesor_id", profesor_id).eq("activo", 1).order("nombre").execute()
    return res.data

def agregar_estudiante(nombre, telefono_acudiente, grado_id, profesor_id):
    supabase.table("estudiantes").insert({
        "nombre": nombre, 
        "telefono_acudiente": telefono_acudiente,
        "grado_id": grado_id, 
        "profesor_id": profesor_id,
        "puntos": 0, "racha": 0, "activo": 1
    }).execute()

def editar_estudiante(estudiante_id, nuevo_nombre, nuevo_telefono):
    supabase.table("estudiantes").update({
        "nombre": nuevo_nombre,
        "telefono_acudiente": nuevo_telefono
    }).eq("id", estudiante_id).execute()

def eliminar_estudiante(estudiante_id):
    supabase.table("estudiantes").update({"activo": 0}).eq("id", estudiante_id).execute()

def ya_registrado_fecha(estudiante_id, fecha_str):
    res = supabase.table("registros").select("id").eq("estudiante_id", estudiante_id).eq("fecha", fecha_str).execute()
    return len(res.data) > 0

def eliminar_asistencia_fecha(estudiante_id, fecha_str):
    supabase.table("registros").delete().eq("estudiante_id", estudiante_id).eq("fecha", fecha_str).execute()

def es_dia_no_lectivo(fecha_str, grado_id, profesor_id):
    try:
        res = supabase.table("dias_no_lectivos").select("*").eq("fecha", fecha_str).eq("grado_id", grado_id).eq("profesor_id", profesor_id).execute()
        return res.data[0] if res.data else None
    except Exception:
        return None

def registrar_dia_no_lectivo(fecha_str, motivo, grado_id, profesor_id):
    try:
        supabase.table("dias_no_lectivos").insert({
            "fecha": fecha_str,
            "motivo": motivo,
            "grado_id": grado_id,
            "profesor_id": profesor_id
        }).execute()
        return True
    except Exception:
        return False

# REGISTRO Y MODIFICACIÓN DE ASISTENCIA
def registrar_asistencia_manual(estudiante, fecha_str, grado_id, profesor_id):
    hora_str = datetime.datetime.now().strftime("%H:%M:%S")
    
    if ya_registrado_fecha(estudiante["id"], fecha_str):
        return None

    res_hoy = supabase.table("registros").select("id", count="exact").eq("fecha", fecha_str).eq("grado_id", grado_id).execute()
    orden = (res_hoy.count or 0) + 1
    
    puntos_extra = PUNTOS_EXTRA_PUNTUALIDAD if orden <= CUPOS_PUNTUALIDAD else 0
    puntos_totales = PUNTOS_BASE + puntos_extra
    
    nuevos_puntos = estudiante["puntos"] + puntos_totales
    
    supabase.table("estudiantes").update({
        "puntos": nuevos_puntos, 
        "ultima_fecha": fecha_str,
        "total_asistencias": estudiante.get("total_asistencias", 0) + 1
    }).eq("id", estudiante["id"]).execute()

    supabase.table("registros").insert({
        "estudiante_id": estudiante["id"], 
        "fecha": fecha_str, 
        "hora": hora_str,
        "puntos_obtenidos": puntos_totales, 
        "orden_llegada": orden,
        "grado_id": grado_id, 
        "profesor_id": profesor_id
    }).execute()

    return {
        "nombre": estudiante["nombre"], 
        "puntos_ganados": puntos_totales,
        "puntos_totales": nuevos_puntos, 
        "racha": estudiante["racha"]
    }

def guardar_excusa(estudiante_id, fecha_str, motivo, profesor_id):
    supabase.table("convivencia").insert({
        "estudiante_id": estudiante_id,
        "fecha": fecha_str,
        "tipo": "Excusa / Justificación",
        "descripcion": f"EXCUSA DE ASISTENCIA ({fecha_str}): {motivo}",
        "profesor_id": profesor_id
    }).execute()

def obtener_excusas_mes(profesor_id, grado_id):
    try:
        res = supabase.table("convivencia").select("*").eq("profesor_id", profesor_id).eq("tipo", "Excusa / Justificación").execute()
        return res.data or []
    except Exception:
        return []

def crear_link_whatsapp(telefono, nombre_estudiante, nombre_profesor):
    if not telefono:
        return None
    num_limpio = re.sub(r'\D', '', str(telefono))
    if len(num_limpio) == 10:
        num_limpio = "57" + num_limpio
        
    mensaje = f"Estimado acudiente, le saluda el/la profe {nombre_profesor} del {NOMBRE_ESCUELA}. Le informamos que el/la estudiante *{nombre_estudiante}* no ha registrado su asistencia a clases en el día de hoy. Por favor confirmar novedades. Muchas gracias."
    msg_encoded = urllib.parse.quote(mensaje)
    return f"https://wa.me/{num_limpio}?text={msg_encoded}"

# ================================================================
# GENERACIÓN DE PDF DE BOLETÍN ACADÉMICO OFICIAL
# ================================================================
def generar_pdf_boletin_oficial(estudiante_nombre, grado_nombre, periodo, sede_nombre, profesor_nombre, dict_notas, observaciones_txt):
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # 1. ESCUDO C.E.R. SIRAVITA
    ruta_escudo = "escudo_siravita.png"
    if os.path.exists(ruta_escudo):
        p.drawImage(ruta_escudo, (width / 2.0) - 25, height - 62, width=50, height=50, mask='auto')

    # 2. ENCABEZADO TEXTO OFICIAL
    p.setFillColor(colors.black)
    p.setFont("Helvetica-Bold", 8.5)
    y_hdr = height - 72
    p.drawCentredString(width / 2.0, y_hdr, "REPUBLICA DE COLOMBIA")
    p.drawCentredString(width / 2.0, y_hdr - 10, "SECRETARIA DE EDUCACION DEPARTAMENTAL NORTE DE SANTANDER")
    p.drawCentredString(width / 2.0, y_hdr - 20, "CENTRO EDUCATIVO RURAL SIRAVITA")
    p.drawCentredString(width / 2.0, y_hdr - 30, "MUNICIPIO DE ARBOLEDAS")

    # 3. BLOQUE VERDE Y AMARILLO INSTITUCIONAL
    y_verde = y_hdr - 62
    p.setFillColor(colors.HexColor("#008037"))
    p.rect(30, y_verde, width - 60, 30, fill=True, stroke=False)
    
    p.setFillColor(colors.white)
    p.setFont("Helvetica-Bold", 7)
    p.drawCentredString(width / 2.0, y_verde + 18, "DANE 254051000139")
    p.drawCentredString(width / 2.0, y_verde + 10, "DECRETO DE CREACIÓN 00252 DEL 12 DE ABRIL DE 2006")

    p.setFillColor(colors.HexColor("#FFF100"))
    p.rect(100, y_verde + 1, width - 200, 8, fill=True, stroke=False)
    p.setFillColor(colors.black)
    p.setFont("Helvetica-Bold", 6)
    p.drawCentredString(width / 2.0, y_verde + 3, "RESOLUCION DE APROBACION DE ESTUDIO 8: 008708-24-10-2024")

    # 4. BANDERA DE COLOMBIA
    y_bandera = y_verde - 42
    ruta_bandera = "bandera_colombia.png"
    if os.path.exists(ruta_bandera):
        p.drawImage(ruta_bandera, (width / 2.0) - 25, y_bandera, width=50, height=38, mask='auto')

    # 5. TÍTULO DEL BOLETÍN
    y_titulo = y_bandera - 16
    p.setFillColor(colors.black)
    p.setFont("Helvetica", 8.5)
    p.drawCentredString(width / 2.0, y_titulo, "Boletín Académico Escuela Nueva de Básica Primaria")

    # 6. DATOS ENCABEZADO
    p.setFont("Helvetica", 8)
    y_d1 = y_titulo - 20
    p.drawString(40, y_d1, f"Sede: {sede_nombre.capitalize()}.........................................................")
    p.drawString(200, y_d1, f"Periodo...{periodo.capitalize()}..................")
    p.drawString(330, y_d1, "año 2026")

    y_d2 = y_d1 - 16
    p.drawString(40, y_d2, f"Estudiante: {estudiante_nombre.upper()}...................................................................................")
    p.drawString(390, y_d2, f"Grado {grado_nombre.capitalize()}............")

    y_d3 = y_d2 - 16
    p.drawString(40, y_d3, f"Docente: {profesor_nombre.title()}............................................................")

    # 7. TABLA DE CALIFICACIONES
    y_tb_top = y_d3 - 18
    h_th = 28
    h_tr = 16.5

    p.setFillColor(colors.HexColor("#F9EBEA"))
    p.rect(30, y_tb_top - h_th, 155, h_th, fill=True, stroke=True)
    
    p.setFillColor(colors.HexColor("#00A2E8"))
    p.rect(185, y_tb_top - h_th, 105, h_th, fill=True, stroke=True)
    p.rect(290, y_tb_top - h_th, 105, h_th, fill=True, stroke=True)
    
    p.setFillColor(colors.HexColor("#FFC90E"))
    p.rect(395, y_tb_top - h_th, 187, h_th, fill=True, stroke=True)

    p.setFillColor(colors.black)
    p.setFont("Helvetica-BoldOblique", 7)
    p.drawString(34, y_tb_top - 16, "AREAS/ ASIGNATURAS")
    p.drawString(188, y_tb_top - 16, "P.E.N. Valoración Cualitativa")
    
    p.setFillColor(colors.blue)
    p.setFont("Helvetica-BoldOblique", 7)
    p.drawString(294, y_tb_top - 16, "Valoración Cuantitativa")
    p.drawString(399, y_tb_top - 16, "Escala Nacional")

    p.setFillColor(colors.black)
    y_row = y_tb_top - h_th

    for mat in MATERIAS_LISTA:
        y_row -= h_tr
        
        key_norm = normalizar_texto(mat)
        nota_val = 0.0
        for k_materia, val_n in dict_notas.items():
            if normalizar_texto(k_materia) == key_norm:
                nota_val = float(val_n)
                break

        val_pen, val_nac = obtener_valoracion_cualitativa(nota_val)

        p.rect(30, y_row, 155, h_tr, fill=False, stroke=True)
        p.setFont("Helvetica-Oblique", 6.5)
        p.drawString(34, y_row + 4, mat[:42])

        p.rect(185, y_row, 105, h_tr, fill=False, stroke=True)
        p.setFont("Helvetica-Oblique", 6.8)
        p.drawString(188, y_row + 4, val_pen)

        p.rect(290, y_row, 105, h_tr, fill=False, stroke=True)
        p.setFont("Helvetica", 7)
        p.drawString(294, y_row + 4, f"{nota_val:.1f}" if nota_val > 0 else "---")

        p.rect(395, y_row, 187, h_tr, fill=False, stroke=True)
        p.setFont("Helvetica-Oblique", 6.8)
        p.drawString(399, y_row + 4, val_nac)

    # 8. OBSERVACIONES
    y_obs_sec = y_row - 22
    p.setFont("Helvetica-Bold", 8)
    p.drawString(60, y_obs_sec, "OBSERVACIONES:")

    p.setFont("Helvetica-Oblique", 7.5)
    texto_obs_m = f"..... {observaciones_txt if observaciones_txt else 'buen rendimiento.'}"
    p.drawString(80, y_obs_sec - 14, texto_obs_m)
    p.drawString(60, y_obs_sec - 16, "." * 115)
    p.drawString(60, y_obs_sec - 28, "." * 115)

    p.showPage()

    # HOJA 2: ESCALA Y FIRMAS
    y_e = height - 120
    p.setFont("Helvetica-BoldOblique", 11)
    p.drawCentredString(width / 2.0, y_e + 20, "ESCALA DE VALORACION INSTITUCIONAL")

    p.setFillColor(colors.HexColor("#00A2E8"))
    p.rect(130, y_e - 30, 130, 30, fill=True, stroke=True)
    p.setFillColor(colors.HexColor("#FFC90E"))
    p.rect(260, y_e - 30, 110, 30, fill=True, stroke=True)
    p.setFillColor(colors.HexColor("#FFF100"))
    p.rect(370, y_e - 30, 110, 30, fill=True, stroke=True)

    p.setFillColor(colors.black)
    p.setFont("Helvetica-Bold", 7)
    p.drawString(135, y_e - 12, "ESCUELA NUEVA")
    p.drawString(135, y_e - 22, "VALOR CUALITATIVA")

    p.drawString(265, y_e - 12, "ESCALA INSTITUCIONAL")
    p.drawString(265, y_e - 22, "VALOR CUALITATIVA")

    p.drawString(375, y_e - 12, "VALORACION")
    p.drawString(375, y_e - 22, "CUANTITATIVA")

    escala_datos = [
        ("EXCELENTE", "DESEMPEÑO SUPERIOR", "4.8 A 5.0"),
        ("BUENO", "DESEMPEÑO ALTO", "4.0 A 4.79"),
        ("ACEPTABLE", "DESEMPEÑO BASICO", "3.0 A 3.99"),
        ("INSUFICIENTE", "DESEMPEÑO BAJO", "1.0 A 2.99")
    ]

    y_f = y_e - 30
    for e_pen, e_inst, e_cuant in escala_datos:
        y_f -= 20
        p.setFont("Helvetica", 8)
        p.rect(130, y_f, 130, 20, fill=False, stroke=True)
        p.drawString(135, y_f + 6, e_pen)

        p.rect(260, y_f, 110, 20, fill=False, stroke=True)
        p.drawString(265, y_f + 6, e_inst)

        p.rect(370, y_f, 110, 20, fill=False, stroke=True)
        p.drawString(375, y_f + 6, e_cuant)

    y_firmas = y_f - 180
    p.setStrokeColor(colors.gray)
    p.line(90, y_firmas, 270, y_firmas)
    p.setFont("Helvetica-Bold", 8.5)
    p.drawCentredString(180, y_firmas - 14, "Docente Titular")

    p.line(340, y_firmas, 520, y_firmas)
    p.drawCentredString(430, y_firmas - 14, "DIRECTOR")

    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer

# GENERACIÓN DE PDF ASISTENCIA
def generar_pdf_asistencia_oficial(grado_nombre, profesor_nombre, registros_mes, excusas_mes, estudiantes_lista, mes_nombre, sede_nombre=SEDE_DEFECTO):
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    ruta_imagen = "plantilla_asistencia.png"
    if os.path.exists(ruta_imagen):
        p.drawImage(ruta_imagen, 0, 0, width=width, height=height)
    else:
        p.setFont("Helvetica-Bold", 10)
        p.drawString(40, height - 40, "CENTRO EDUCATIVO RURAL SIRAVITA - CONTROL ASISTENCIA")

    p.setFont("Helvetica-Bold", 9.5)
    p.setFillColor(colors.HexColor("#111827"))
    p.drawString(110, 415, str(sede_nombre).upper())
    p.drawString(330, 415, str(mes_nombre).upper())

    set_excusas = set()
    for exc in excusas_mes:
        try:
            e_id = exc["estudiante_id"]
            f_str = exc["fecha"]
            dia_num = int(f_str.split("-")[2])
            set_excusas.add((e_id, dia_num))
        except Exception:
            pass

    x_alumnos = 8
    x_grado = 78
    x_dias_inicio = 138.5
    w_dia = 13.5
    x_total = 560

    y_fila_inicio = 328
    h_fila = 19.1

    for idx, est in enumerate(estudiantes_lista[:12]):
        y_pos = y_fila_inicio - (idx * h_fila)
        
        p.setFont("Helvetica-Bold", 6.8)
        p.setFillColor(colors.HexColor("#111827"))
        p.drawString(x_alumnos, y_pos, str(est["nombre"])[:18])

        p.setFont("Helvetica-Bold", 6.8)
        p.setFillColor(colors.HexColor("#374151"))
        p.drawString(x_grado, y_pos, str(grado_nombre)[:6])

        tot_asist = 0
        for d in range(1, 32):
            asistio = any(r["estudiante_id"] == est["id"] and int(r["fecha"].split("-")[2]) == d for r in registros_mes)
            tiene_excusa = (est["id"], d) in set_excusas
            x_check = x_dias_inicio + ((d - 1) * w_dia)
            
            if asistio:
                tot_asist += 1
                p.setFont("Helvetica-Bold", 11)
                p.setFillColor(colors.HexColor("#1B5E20"))
                p.drawString(x_check - 1, y_pos - 1, "✓")
            elif tiene_excusa:
                p.setFont("Helvetica-Bold", 8.5)
                p.setFillColor(colors.HexColor("#E65100"))
                p.drawString(x_check + 1, y_pos, "E")

        p.setFont("Helvetica-Bold", 9.5)
        p.setFillColor(colors.HexColor("#111827"))
        p.drawString(x_total, y_pos, str(tot_asist))

    p.setFont("Helvetica-Bold", 9.5)
    p.setFillColor(colors.HexColor("#111827"))
    p.drawString(115, 28, str(profesor_nombre).upper())

    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer

# GENERACIÓN EXCEL
def generar_excel_asistencia_oficial(grado_nombre, profesor_nombre, registros_mes, excusas_mes, estudiantes_lista, mes_nombre, sede_nombre=SEDE_DEFECTO):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Control Asistencia"
    ws.views.sheetView[0].showGridLines = True

    VERDE_OSCURO = "004D25"
    VERDE_LEMA = "2E7D32"
    VERDE_HEADER_TABLA = "1B432C"
    AMARILLO_CREMA = "FFF9C4"
    
    fill_header_tabla = PatternFill(start_color=VERDE_HEADER_TABLA, end_color=VERDE_HEADER_TABLA, fill_type="solid")
    fill_row_top = PatternFill(start_color=AMARILLO_CREMA, end_color=AMARILLO_CREMA, fill_type="solid")

    font_encabezado_bold = Font(name="Calibri", size=10, bold=True, color="000000")
    font_lema = Font(name="Calibri", size=10, bold=True, italic=True, color=VERDE_LEMA)
    font_subtitulos = Font(name="Calibri", size=11, bold=True, color="000000")
    font_th = Font(name="Calibri", size=9, bold=True, color="FFFFFF")
    font_data = Font(name="Calibri", size=9)

    align_center = Alignment(horizontal="center", vertical="center")
    align_left = Alignment(horizontal="left", vertical="center")

    thin_border = Border(
        left=Side(style='thin', color='B0BEC5'),
        right=Side(style='thin', color='B0BEC5'),
        top=Side(style='thin', color='B0BEC5'),
        bottom=Side(style='thin', color='B0BEC5')
    )

    set_excusas = set()
    for exc in excusas_mes:
        try:
            e_id = exc["estudiante_id"]
            f_str = exc["fecha"]
            dia_num = int(f_str.split("-")[2])
            set_excusas.add((e_id, dia_num))
        except Exception:
            pass

    ws.merge_cells("A1:AH1")
    ws["A1"] = "REPÚBLICA DE COLOMBIA"
    ws["A1"].font = font_encabezado_bold
    ws["A1"].alignment = align_center

    ws.merge_cells("A2:AH2")
    ws["A2"] = "SECRETARÍA DE EDUCACIÓN DEPARTAMENTAL NORTE DE SANTANDER"
    ws["A2"].font = font_encabezado_bold
    ws["A2"].alignment = align_center

    ws.merge_cells("A3:AH3")
    ws["A3"] = "CENTRO EDUCATIVO RURAL SIRAVITA"
    ws["A3"].font = font_encabezado_bold
    ws["A3"].alignment = align_center

    ws.merge_cells("A4:AH4")
    ws["A4"] = "MUNICIPIO DE ARBOLEDAS"
    ws["A4"].font = font_encabezado_bold
    ws["A4"].alignment = align_center

    ws.merge_cells("A5:AH5")
    ws["A5"] = "DANE 254051000139 | DECRETO DE CREACIÓN 00252 DEL 12 DE ABRIL DE 2005"
    ws["A5"].font = Font(name="Calibri", size=8)
    ws["A5"].alignment = align_center

    ws.merge_cells("A6:AH6")
    ws["A6"] = "Lema: Con Escuela nueva y metodología activa para una educación proactiva."
    ws["A6"].font = font_lema
    ws["A6"].alignment = align_center

    ws.merge_cells("A8:AH8")
    ws["A8"] = "REGISTRO CONTROL ASISTENCIA"
    ws["A8"].font = Font(name="Calibri", size=14, bold=True, color=VERDE_OSCURO)
    ws["A8"].alignment = align_center

    ws.merge_cells("A10:AH10")
    ws["A10"] = f"SEDE: {sede_nombre.upper()}     |     MES DE: {mes_nombre.upper()}     |     AÑO: 2026"
    ws["A10"].font = font_subtitulos
    ws["A10"].alignment = align_left

    headers = ["ALUMNOS", "GRADO"] + [str(i) for i in range(1, 32)] + ["TOTAL"]
    
    for col_idx, text in enumerate(headers, start=1):
        cell = ws.cell(row=12, column=col_idx, value=text)
        cell.font = font_th
        cell.fill = fill_header_tabla
        cell.alignment = align_center
        cell.border = thin_border

    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=13, column=col_idx)
        cell.fill = fill_row_top
        cell.border = thin_border

    row_start = 14
    for est in estudiantes_lista:
        ws.cell(row=row_start, column=1, value=est["nombre"]).font = font_data
        ws.cell(row=row_start, column=1).alignment = align_left
        ws.cell(row=row_start, column=1).border = thin_border

        ws.cell(row=row_start, column=2, value=grado_nombre).font = font_data
        ws.cell(row=row_start, column=2).alignment = align_center
        ws.cell(row=row_start, column=2).border = thin_border

        asistencias_total = 0
        for dia in range(1, 32):
            col_idx = dia + 2
            asistio = any(r["estudiante_id"] == est["id"] and int(r["fecha"].split("-")[2]) == dia for r in registros_mes)
            tiene_excusa = (est["id"], dia) in set_excusas
            
            if asistio:
                val = "✓"
                color_txt = "2E7D32"
                asistencias_total += 1
            elif tiene_excusa:
                val = "E"
                color_txt = "E65100"
            else:
                val = ""
                color_txt = "000000"

            cell_d = ws.cell(row=row_start, column=col_idx, value=val)
            cell_d.font = Font(name="Calibri", size=10, bold=True, color=color_txt)
            cell_d.alignment = align_center
            cell_d.border = thin_border

        cell_t = ws.cell(row=row_start, column=34, value=asistencias_total)
        cell_t.font = Font(name="Calibri", size=9, bold=True)
        cell_t.alignment = align_center
        cell_t.border = thin_border

        row_start += 1

    while row_start < 28:
        for col_idx in range(1, 35):
            c = ws.cell(row=row_start, column=col_idx)
            c.border = thin_border
        row_start += 1

    ws.cell(row=30, column=1, value=f"DOCENTE: {profesor_nombre.upper()}").font = font_subtitulos
    ws.cell(row=32, column=1, value="Documento Oficial - Uso Académico (✓ = Asistencia, E = Excusa)").font = Font(name="Calibri", size=8, italic=True)

    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 12
    for col in range(3, 34):
        col_letter = get_column_letter(col)
        ws.column_dimensions[col_letter].width = 3.5
    ws.column_dimensions['AH'].width = 8

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

# LÓGICA ADAPTADOR DE GUÍAS
def extraer_texto_archivo(archivo_subido, tipo_archivo):
    texto = ""
    try:
        if tipo_archivo == "pdf":
            reader = PdfReader(archivo_subido)
            for page in reader.pages:
                t = page.extract_text()
                if t: texto += t + "\n"
        elif tipo_archivo == "docx":
            doc = docx.Document(archivo_subido)
            for p in doc.paragraphs:
                texto += p.text + "\n"
        elif tipo_archivo == "txt":
            texto = archivo_subido.getvalue().decode("utf-8")
    except Exception as e:
        st.error(f"Error al leer el archivo: {e}")
    return texto.strip()

def adaptar_contenido_escuela_nueva(texto_original, eje_tematico):
    lineas = [l.strip() for l in texto_original.split("\n") if l.strip()]
    total = len(lineas)
    
    if total == 0:
        lineas = ["Contenido genérico sobre el eje temático: " + eje_tematico]
        total = 1

    q1, q2, q3 = total // 4, total // 2, (3 * total) // 4
    
    momento_a = "\n".join(lineas[:q1]) if q1 > 0 else "• Responde en tu cuaderno: ¿Qué sabes acerca de " + eje_tematico + "?\n• Observa las situaciones cotidianas relacionadas con el tema y comparte tus ideas."
    momento_b = "\n".join(lineas[q1:q2]) if q2 > q1 else "• Concepto clave: " + eje_tematico + ".\n• Lee atentamente la explicación de la guía y consigna los puntos principales en tu cuaderno de trabajo."
    momento_c = "\n".join(lineas[q2:q3]) if q3 > q2 else "• Desarrolla las actividades de ejercitación individuales y en parejas.\n• Resuelve los problemas prácticos propuestos para afianzar el conocimiento."
    momento_d = "\n".join(lineas[q3:]) if total > q3 else "• Demuestra lo aprendido resolviendo la evaluación final.\n• Aplica lo aprendido con ayuda de tu familia en la casa."

    return {
        "A": momento_a,
        "B": momento_b,
        "C": momento_c,
        "D": momento_d
    }

def generar_pdf_guia_adaptada(materia, eje_tematico, momentos, profesor_nombre, grado_nombre):
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    p.setFont("Helvetica-Bold", 9)
    p.drawCentredString(width / 2.0, height - 35, "CENTRO EDUCATIVO RURAL SIRAVITA - ARBOLEDAS")
    p.setFont("Helvetica-Bold", 11)
    p.setFillColor(colors.HexColor("#008037"))
    p.drawCentredString(width / 2.0, height - 50, f"GUÍA DE APRENDIZAJE DE ESCUELA NUEVA — {materia.upper()}")
    
    p.setFillColor(colors.black)
    p.setFont("Helvetica", 8.5)
    p.drawString(40, height - 70, f"Grado: {grado_nombre}   |   Docente: {profesor_nombre}   |   Eje Temático: {eje_tematico}")
    p.line(40, height - 75, width - 40, height - 75)

    y_pos = height - 95
    secciones = [
        ("MOMENTO A: Vivencia / Saberes Previos", momentos["A"], "#1B432C"),
        ("MOMENTO B: Fundamentación Teórica", momentos["B"], "#00A2E8"),
        ("MOMENTO C: Ejercitación y Aplicación", momentos["C"], "#FFC90E"),
        ("MOMENTO D: Evaluación y Extensión", momentos["D"], "#E65100")
    ]

    for tit, cont, col_hex in secciones:
        if y_pos < 100:
            p.showPage()
            y_pos = height - 50

        p.setFillColor(colors.HexColor(col_hex))
        p.rect(40, y_pos - 18, width - 80, 18, fill=True, stroke=False)
        p.setFillColor(colors.white if col_hex != "#FFC90E" else colors.black)
        p.setFont("Helvetica-Bold", 9)
        p.drawString(48, y_pos - 13, tit)

        y_pos -= 30
        p.setFillColor(colors.black)
        p.setFont("Helvetica", 8)
        
        for lin in cont.split("\n")[:12]:
            if y_pos < 50:
                p.showPage()
                y_pos = height - 50
            p.drawString(48, y_pos, lin[:100])
            y_pos -= 12
        y_pos -= 10

    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer

def generar_docx_guia_adaptada(materia, eje_tematico, momentos, profesor_nombre, grado_nombre):
    doc = docx.Document()
    
    doc.add_heading('CENTRO EDUCATIVO RURAL SIRAVITA', level=1)
    p = doc.add_paragraph()
    p.add_run(f"GUÍA DE APRENDIZAJE ESCUELA NUEVA — {materia.upper()}\n").bold = True
    p.add_run(f"Grado: {grado_nombre} | Docente: {profesor_nombre} | Eje Temático: {eje_tematico}\n")

    secciones = [
        ("MOMENTO A: Vivencia / Saberes Previos", momentos["A"]),
        ("MOMENTO B: Fundamentación Teórica", momentos["B"]),
        ("MOMENTO C: Ejercitación y Aplicación", momentos["C"]),
        ("MOMENTO D: Evaluación y Extensión", momentos["D"])
    ]

    for tit, cont in secciones:
        doc.add_heading(tit, level=2)
        doc.add_paragraph(cont)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_html_guia_adaptada(materia, eje_tematico, momentos, profesor_nombre, grado_nombre):
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>Guía {materia}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 30px; line-height: 1.6; }}
            .header {{ text-align: center; border-bottom: 2px solid #008037; padding-bottom: 10px; }}
            .momento {{ background: #f4f4f4; padding: 10px; border-left: 5px solid #008037; margin-top: 15px; font-weight: bold; }}
            .contenido {{ padding: 10px; background: #fff; border: 1px solid #ddd; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h2>CENTRO EDUCATIVO RURAL SIRAVITA</h2>
            <h3>GUÍA DE APRENDIZAJE ESCUELA NUEVA — {materia.upper()}</h3>
            <p><strong>Grado:</strong> {grado_nombre} | <strong>Docente:</strong> {profesor_nombre} | <strong>Eje:</strong> {eje_tematico}</p>
        </div>
        <div class="momento">MOMENTO A: Vivencia / Saberes Previos</div>
        <div class="contenido"><pre>{momentos['A']}</pre></div>
        <div class="momento">MOMENTO B: Fundamentación Teórica</div>
        <div class="contenido"><pre>{momentos['B']}</pre></div>
        <div class="momento">MOMENTO C: Ejercitación y Aplicación</div>
        <div class="contenido"><pre>{momentos['C']}</pre></div>
        <div class="momento">MOMENTO D: Evaluación y Extensión</div>
        <div class="contenido"><pre>{momentos['D']}</pre></div>
    </body>
    </html>
    """

# ================================================================
# VENTANA CELEBRACIÓN
# ================================================================
@st.dialog("🎉 ¡Asistencia Registrada!", width="large")
def ventana_celebracion(res):
    nombre = res["nombre"]
    puntos_ganados = res["puntos_ganados"]
    puntos_totales = res["puntos_totales"]
    racha = res["racha"]
    
    saludo_hablado = "Buenos días"
    frase_animo = random.choice(MENSAJES_ANIMO)
    texto_voz = f"Buenos días, bienvenido {nombre}. Has ganado {puntos_ganados} puntos."

    st.markdown(f"""
        <div style="background-color: #22C55E; color: white; padding: 25px; border-radius: 20px; text-align: center;">
            <h2 style="color: #FFE066; margin:0;">🌟 ¡{saludo_hablado}, {nombre}! 🌟</h2>
            <h1 style="color: #FFE066; font-size: 40px; margin: 10px 0;">{nombre}</h1>
            <p style="font-size: 20px; margin:5px;">✨ +{puntos_ganados} puntos ganados</p>
            <p style="font-size: 22px; font-weight: bold; margin:5px;">🔥 Racha: {racha} día(s) | Puntos: {puntos_totales}</p>
            <h3 style="color: #FFE066; margin-top:10px;">{frase_animo}</h3>
        </div>
    """, unsafe_allow_html=True)

    components.html(f"""
        <script>
            (function() {{
                if ('speechSynthesis' in window) {{
                    window.speechSynthesis.cancel();
                    var msg = new SpeechSynthesisUtterance("{texto_voz}");
                    msg.lang = 'es-ES';
                    msg.rate = 1.1;
                    window.speechSynthesis.speak(msg);
                }}
            }})();
        </script>
    """, height=0)

    st.write("")
    if st.button("☑️ ¡Entendido!", use_container_width=True, type="primary"):
        st.rerun()

# ================================================================
# VISTA: LOGIN / REGISTRO
# ================================================================
if st.session_state.profesor is None:
    st.title("🏫 " + NOMBRE_ESCUELA)
    st.subheader("Plataforma Multidocente de Gestión Educativa")
    
    tab_login, tab_registro = st.tabs(["🔑 Iniciar Sesión", "📝 Registrar nuevo Profesor"])
    
    with tab_login:
        with st.form("form_login"):
            email = st.text_input("Correo electrónico:")
            password = st.text_input("Contraseña:", type="password")
            btn_login = st.form_submit_button("Ingresar al Panel")
            if btn_login:
                prof = login_profesor(email, password)
                if prof:
                    st.session_state.profesor = prof
                    st.success(f"¡Bienvenido(a), Profe {prof['nombre']}!")
                    st.rerun()
                else:
                    st.error("Correo o contraseña incorrectos.")

    with tab_registro:
        with st.form("form_reg"):
            nom_reg = st.text_input("Nombre Completo:")
            email_reg = st.text_input("Correo electrónico:")
            pass_reg = st.text_input("Crear Contraseña:", type="password")
            btn_reg = st.form_submit_button("Crear Cuenta de Docente")
            if btn_reg:
                if nom_reg and email_reg and pass_reg:
                    ok, msg = registrar_profesor(nom_reg, email_reg, pass_reg)
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)
                else:
                    st.warning("Completa todos los campos.")

# ================================================================
# VISTA: PANEL PRINCIPAL
# ================================================================
else:
    prof = st.session_state.profesor
    
    with st.sidebar:
        st.title(f"👨‍🏫 Profe: {prof['nombre']}")
        st.caption(f"📧 {prof['email']}")
        
        if st.button("🚪 Cerrar Sesión", use_container_width=True):
            st.session_state.profesor = None
            st.rerun()
            
        st.markdown("---")
        
        grados = obtener_grados(prof["id"])
        if not grados:
            st.warning("Aún no tienes cursos creados.")
            grado_sel_id = None
            grado_sel_nombre = ""
        else:
            nombres_grados = [g["nombre"] for g in grados]
            grado_sel_nombre = st.selectbox("Selecciona un Grado:", nombres_grados)
            grado_sel_id = next(g["id"] for g in grados if g["nombre"] == grado_sel_nombre)

        st.markdown("---")
        
        with st.expander("⚙️ Gestionar Cursos / Grados"):
            nuevo_grado_txt = st.text_input("Nombre de nuevo curso:", placeholder="Ej. Grado Cuarto")
            if st.button("➕ Crear Curso", use_container_width=True):
                if nuevo_grado_txt.strip():
                    agregar_grado(nuevo_grado_txt.strip(), prof["id"])
                    st.success("Curso creado.")
                    st.rerun()
            
            if grado_sel_id:
                st.markdown("---")
                edit_g_nom = st.text_input("Editar nombre del curso actual:", value=grado_sel_nombre)
                if st.button("✏️ Guardar Nombre Curso", use_container_width=True):
                    editar_grado(grado_sel_id, edit_g_nom.strip())
                    st.success("Nombre actualizado.")
                    st.rerun()
                    
                if st.button("🗑️ Eliminar Curso Actual", type="primary", use_container_width=True):
                    eliminar_grado(grado_sel_id)
                    st.success("Curso eliminado.")
                    st.rerun()

        st.markdown("---")
        with st.expander("👤 Agregar / Editar Estudiante"):
            if grado_sel_id:
                nom_est = st.text_input("Nombre completo de alumno:")
                tel_est = st.text_input("Teléfono WhatsApp Acudiente:", placeholder="Ej: 3001234567")
                if st.button("➕ Guardar Alumno", use_container_width=True):
                    if nom_est.strip():
                        agregar_estudiante(nom_est.strip(), tel_est.strip(), grado_sel_id, prof["id"])
                        st.success("Alumno guardado.")
                        st.rerun()
                
                st.markdown("---")
                estudiantes_lista = obtener_estudiantes(grado_sel_id, prof["id"])
                if estudiantes_lista:
                    dict_est = {e["nombre"]: e for e in estudiantes_lista}
                    est_edit_nom = st.selectbox("Selecciona alumno a editar:", list(dict_est.keys()))
                    est_obj = dict_est[est_edit_nom]
                    
                    nuevo_nom_val = st.text_input("Nuevo nombre:", value=est_obj["nombre"])
                    nuevo_tel_val = st.text_input("Nuevo teléfono:", value=est_obj.get("telefono_acudiente") or "")
                    
                    col_e1, col_e2 = st.columns(2)
                    with col_e1:
                        if st.button("✏️ Actualizar", use_container_width=True):
                            editar_estudiante(est_obj["id"], nuevo_nom_val.strip(), nuevo_tel_val.strip())
                            st.success("Datos actualizados.")
                            st.rerun()
                    with col_e2:
                        if st.button("🗑️ Eliminar", use_container_width=True):
                            eliminar_estudiante(est_obj["id"])
                            st.success("Alumno eliminado.")
                            st.rerun()

    # PESTAÑAS PRINCIPALES
    st.title(f"📋 Asistente Educativo — {grado_sel_nombre if grado_sel_nombre else 'Crea un curso'}")

    if grado_sel_id:
        t_docs, t_asistencia, t_estadisticas, t_notas, t_boletines, t_adaptador, t_convivencia, t_lideres = st.tabs([
            "📄 Documentos Institucionales",
            "📋 Registro de Asistencia", 
            "📊 Resumen Estadístico Mensual",
            "📝 Calificaciones Continuas",
            "📄 Boletines Académicos",
            "🧩 Adaptador de Guías",
            "⚖️ Observador de Convivencia",
            "🏆 Tabla de Líderes"
        ])

        # 1. DOCUMENTOS INSTITUCIONALES
        with t_docs:
            st.subheader("📁 Repositorio de Documentos por Materia")
            tipo_doc_sel = st.radio("Categoría:", ["Planes de Área", "Planes de Clase", "Guías Educativas"], horizontal=True)
            materia_doc = st.selectbox("Materia:", MATERIAS_LISTA)
            
            archivo_subido = st.file_uploader(f"Subir documento a {tipo_doc_sel} ({materia_doc}):", type=["pdf", "docx", "pptx", "xlsx", "txt"])
            if archivo_subido is not None:
                if st.button("💾 Guardar Documento en la Nube"):
                    bytes_data = archivo_subido.getvalue()
                    b64_str = base64.b64encode(bytes_data).decode('utf-8')
                    
                    supabase.table("documentos").insert({
                        "nombre": archivo_subido.name,
                        "materia": materia_doc,
                        "tipo_doc": tipo_doc_sel,
                        "contenido_b64": b64_str,
                        "grado_id": grado_sel_id,
                        "profesor_id": prof["id"]
                    }).execute()
                    st.success("¡Documento guardado permanentemente en la nube!")
                    st.rerun()

            st.markdown("---")
            st.write(f"**Documentos de {materia_doc} en {tipo_doc_sel}:**")
            res_docs = supabase.table("documentos").select("*").eq("materia", materia_doc).eq("tipo_doc", tipo_doc_sel).eq("grado_id", grado_sel_id).eq("profesor_id", prof["id"]).execute()
            
            if not res_docs.data:
                st.info(f"Aún no se han subido documentos para {materia_doc}.")
            else:
                for doc in res_docs.data:
                    c1, c2 = st.columns([3, 1])
                    with c1:
                        st.write(f"📄 **{doc['nombre']}** ({doc['created_at'][:10]})")
                    with c2:
                        bytes_dec = base64.b64decode(doc['contenido_b64'])
                        st.download_button("⬇️ Descargar", data=bytes_dec, file_name=doc['nombre'], use_container_width=True)

        # 2. REGISTRO DE ASISTENCIA
        with t_asistencia:
            estudiantes = obtener_estudiantes(grado_sel_id, prof["id"])
            
            col_f1, col_f2 = st.columns([1, 2])
            with col_f1:
                fecha_gestion = st.date_input("📅 Selecciona Fecha a Gestionar:", datetime.date.today())
                fecha_sel_str = fecha_gestion.isoformat()
            
            with col_f2:
                es_fin = fecha_gestion.weekday() >= 5
                reg_no_lect = es_dia_no_lectivo(fecha_sel_str, grado_sel_id, prof["id"])
                if es_fin:
                    st.caption("ℹ️ *Fecha seleccionada corresponde a Fin de Semana.*")
                if reg_no_lect:
                    st.warning(f"⚠️ Día No Lectivo: {reg_no_lect['motivo']}")

            total_est = len(estudiantes) if estudiantes else 0
            asistieron_count = 0
            estudiantes_ausentes = []
            
            if estudiantes:
                for est in estudiantes:
                    if ya_registrado_fecha(est["id"], fecha_sel_str):
                        asistieron_count += 1
                    else:
                        estudiantes_ausentes.append(est)
            
            faltaron_count = total_est - asistieron_count

            c_red, c_green = st.columns(2)
            with c_red:
                st.markdown(f"""
                    <div style="background-color: #FFEBEE; border: 2px solid #EF5350; border-radius: 12px; padding: 8px 12px; text-align: center;">
                        <span style="color: #C62828; font-weight: bold; font-size: 13px;">✖ Faltaron / Sin Registro</span><br>
                        <span style="color: #C62828; font-size: 26px; font-weight: bold;">{faltaron_count}</span>
                    </div>
                """, unsafe_allow_html=True)
            with c_green:
                st.markdown(f"""
                    <div style="background-color: #E8F5E9; border: 2px solid #66BB6A; border-radius: 12px; padding: 8px 12px; text-align: center;">
                        <span style="color: #2E7D32; font-weight: bold; font-size: 13px;">✅ Asistieron</span><br>
                        <span style="color: #2E7D32; font-size: 26px; font-weight: bold;">{asistieron_count}</span>
                    </div>
                """, unsafe_allow_html=True)

            st.write("")

            if estudiantes:
                mes_actual_nombre = MESES_ESPANOL[fecha_gestion.month - 1]
                res_reg = supabase.table("registros").select("*").eq("grado_id", grado_sel_id).execute()
                excusas_list = obtener_excusas_mes(prof["id"], grado_sel_id)
                
                col_ex1, col_ex2 = st.columns(2)
                with col_ex1:
                    pdf_bytes = generar_pdf_asistencia_oficial(
                        grado_nombre=grado_sel_nombre,
                        profesor_nombre=prof['nombre'],
                        registros_mes=res_reg.data or [],
                        excusas_mes=excusas_list,
                        estudiantes_lista=estudiantes,
                        mes_nombre=mes_actual_nombre,
                        sede_nombre=SEDE_DEFECTO
                    )
                    st.download_button(
                        label="📄 Exportar Planilla Oficial en PDF (Con ✓ y E)",
                        data=pdf_bytes,
                        file_name=f"Planilla_Oficial_Asistencia_{grado_sel_nombre}_{mes_actual_nombre}_2026.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        type="primary"
                    )
                with col_ex2:
                    excel_bytes = generar_excel_asistencia_oficial(
                        grado_nombre=grado_sel_nombre,
                        profesor_nombre=prof['nombre'],
                        registros_mes=res_reg.data or [],
                        excusas_mes=excusas_list,
                        estudiantes_lista=estudiantes,
                        mes_nombre=mes_actual_nombre,
                        sede_nombre=SEDE_DEFECTO
                    )
                    st.download_button(
                        label="📊 Exportar Planilla Oficial en Excel (.xlsx)",
                        data=excel_bytes,
                        file_name=f"Control_Asistencia_{grado_sel_nombre}_{mes_actual_nombre}_2026.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )

            st.markdown("---")

            st.subheader(f"📋 Lista de Asistencia — {fecha_sel_str}")
            if not estudiantes:
                st.info("No hay alumnos registrados en este curso.")
            else:
                cols = st.columns(3)
                for i, est in enumerate(estudiantes):
                    registrado = ya_registrado_fecha(est["id"], fecha_sel_str)
                    col = cols[i % 3]
                    with col:
                        with st.container(border=True):
                            if registrado:
                                st.success(f"✅ {est['nombre']}")
                                if st.button("❌ Quitar Asistencia", key=f"del_{est['id']}_{fecha_sel_str}", use_container_width=True):
                                    eliminar_asistencia_fecha(est["id"], fecha_sel_str)
                                    st.toast(f"Asistencia removida para {est['nombre']}")
                                    st.rerun()
                            else:
                                st.markdown(f"**👤 {est['nombre']}**")
                                if st.button("✋ Marcar Asistencia", key=f"add_{est['id']}_{fecha_sel_str}", use_container_width=True):
                                    res = registrar_asistencia_manual(est, fecha_sel_str, grado_sel_id, prof["id"])
                                    if res and fecha_sel_str == datetime.date.today().isoformat():
                                        st.balloons()
                                        ventana_celebracion(res)
                                    else:
                                        st.rerun()

            st.markdown("---")

            with st.expander("📝 Registrar Excusa o Justificación de Inasistencia", expanded=False):
                if estudiantes:
                    dict_e_ex = {e["nombre"]: e["id"] for e in estudiantes}
                    est_excusa_nom = st.selectbox("Selecciona Estudiante para Excusa:", list(dict_e_ex.keys()))
                    motivo_excusa = st.text_area("Motivo o detalle de la excusa médica/permiso:", placeholder="Ej. Presenta excusa médica por cuadro febril.")
                    if st.button("💾 Guardar Excusa (Aparecerá como 'E' en el PDF)", use_container_width=True):
                        if motivo_excusa.strip():
                            guardar_excusa(dict_e_ex[est_excusa_nom], fecha_sel_str, motivo_excusa.strip(), prof["id"])
                            st.success(f"Excusa registrada correctamente para {est_excusa_nom}.")
                            st.rerun()

            with st.expander("🚫 Declarar Fecha como Día Sin Clase (Festivo/Paro/Jornada)", expanded=False):
                motivo_txt = st.text_input("Motivo de suspensión:", placeholder="Ej. Jornada Pedagógica")
                if st.button("📌 Guardar Día Sin Clase"):
                    if motivo_txt.strip():
                        registrar_dia_no_lectivo(fecha_sel_str, motivo_txt.strip(), grado_sel_id, prof["id"])
                        st.success("Día guardado.")
                        st.rerun()

            if estudiantes_ausentes:
                with st.expander("📲 Notificar ausencias por WhatsApp", expanded=False):
                    for aus in estudiantes_ausentes:
                        tel = aus.get("telefono_acudiente")
                        c_a1, c_a2 = st.columns([2, 1])
                        with c_a1:
                            st.write(f"🔴 **{aus['nombre']}**")
                        with c_a2:
                            link_wa = crear_link_whatsapp(tel, aus['nombre'], prof['nombre'])
                            if link_wa:
                                st.link_button("📲 WhatsApp", link_wa, use_container_width=True)
                            else:
                                st.caption("Sin teléfono")

        # 3. RESUMEN ESTADÍSTICO MENSUAL POR GRADO
        with t_estadisticas:
            st.subheader(f"📊 Dashboard Estadístico y Métricas Mensuales — {grado_sel_nombre}")
            estudiantes = obtener_estudiantes(grado_sel_id, prof["id"])
            
            if not estudiantes:
                st.info("Agrega estudiantes para visualizar el dashboard estadístico.")
            else:
                res_reg = supabase.table("registros").select("*").eq("grado_id", grado_sel_id).execute()
                data_reg = res_reg.data or []
                excusas_list = obtener_excusas_mes(prof["id"], grado_sel_id)
                
                tot_estudiantes = len(estudiantes)
                tot_asistencias_mes = len(data_reg)
                tot_excusas_mes = len(excusas_list)
                
                fechas_registradas = set(r["fecha"] for r in data_reg) if data_reg else set()
                dias_clase_contados = max(len(fechas_registradas), 1)
                
                asistencias_esperadas = tot_estudiantes * dias_clase_contados
                pct_asistencia = round((tot_asistencias_mes / asistencias_esperadas) * 100, 1) if asistencias_esperadas > 0 else 0
                
                m1, m2, m3, m4, m5 = st.columns(5)
                with m1:
                    st.metric("👥 Estudiantes Activos", tot_estudiantes)
                with m2:
                    st.metric("📅 Días Lectivos Registrados", dias_clase_contados)
                with m3:
                    st.metric("✅ % Asistencia del Mes", f"{pct_asistencia}%", delta=f"{pct_asistencia - 80:.1f}% vs Meta 80%")
                with m4:
                    st.metric("📝 Excusas Justificadas", tot_excusas_mes)
                with m5:
                    prom_puntos = int(sum(e.get("puntos", 0) for e in estudiantes) / tot_estudiantes) if tot_estudiantes else 0
                    st.metric("⭐ Puntos Promedio", prom_puntos)

                st.markdown("---")
                
                g1, g2 = st.columns([1.8, 1.2])
                
                with g1:
                    st.write("**📈 Distribución Diaria de Asistencia:**")
                    if data_reg:
                        df_reg = pd.DataFrame(data_reg)
                        df_reg["fecha_dt"] = pd.to_datetime(df_reg["fecha"])
                        df_diario = df_reg.groupby(df_reg["fecha_dt"].dt.strftime("%d-%b")).size().reset_index(name="Asistencias")
                        st.bar_chart(df_diario.set_index("fecha_dt")["Asistencias"], use_container_width=True)
                    else:
                        st.info("Aún no hay asistencias tomadas este mes para graficar la tendencia diaria.")

                with g2:
                    st.write("**🏆 Alumno Destacado y Alertas:**")
                    conteo_ind = {}
                    for r in data_reg:
                        eid = r["estudiante_id"]
                        conteo_ind[eid] = conteo_ind.get(eid, 0) + 1
                    
                    id_a_nombre = {e["id"]: e["nombre"] for e in estudiantes}
                    
                    if conteo_ind:
                        top_id = max(conteo_ind, key=conteo_ind.get)
                        st.success(f"🥇 **Mayor Asistencia:** {id_a_nombre.get(top_id, 'N/A')} ({conteo_ind[top_id]} asistencias)")
                    
                    baja_asist = [e["nombre"] for e in estudiantes if conteo_ind.get(e["id"], 0) < (dias_clase_contados * 0.7)]
                    if baja_asist:
                        st.warning(f"⚠️ **Alerta Inasistencia (>30% faltas):**\n" + "\n".join([f"• {nom}" for nom in baja_asist]))
                    else:
                        st.info("🎉 ¡Excelente! Ningún estudiante presenta faltas críticas.")

                st.markdown("---")
                
                st.write("**📊 Promedio General de Calificaciones por Asignatura:**")
                res_notas = supabase.table("calificaciones").select("materia, nota").eq("profesor_id", prof["id"]).execute()
                if res_notas.data:
                    df_notas_g = pd.DataFrame(res_notas.data)
                    df_prom_mat = df_notas_g.groupby("materia")["nota"].mean().reset_index()
                    df_prom_mat["nota"] = df_prom_mat["nota"].round(2)
                    st.bar_chart(df_prom_mat.set_index("materia")["nota"], use_container_width=True)
                else:
                    st.caption("Ingresa calificaciones en la pestaña 'Calificaciones' para visualizar los promedios por materia.")

        # 4. CALIFICACIONES CONTINUAS (CON NOMBRE Y FECHA DE ACTIVIDAD)
        with t_notas:
            st.subheader(f"📝 Evaluación Continua — {grado_sel_nombre}")
            
            estudiantes = obtener_estudiantes(grado_sel_id, prof["id"])
            
            if not estudiantes:
                st.info(f"No hay alumnos registrados en **{grado_sel_nombre}**. Agrega estudiantes desde el menú lateral.")
            else:
                dict_e = {e["nombre"]: e for e in estudiantes}
                
                col_e1, col_e2, col_e3 = st.columns(3)
                with col_e1:
                    est_sel_nombre = st.selectbox(f"Estudiante ({grado_sel_nombre}):", list(dict_e.keys()))
                    est_sel_obj = dict_e[est_sel_nombre]
                with col_e2:
                    mat_sel = st.selectbox("Materia / Asignatura:", MATERIAS_LISTA)
                with col_e3:
                    per_sel = st.selectbox("Periodo Académico:", ["Primero", "Segundo", "Tercero", "Cuarto"])

                st.markdown("---")
                
                c_head_izq, c_head_der = st.columns([1.6, 1.4])
                
                with c_head_izq:
                    st.write(f"### ✏️ Planilla de **{est_sel_nombre}**\n*{grado_sel_nombre} | {mat_sel} — Periodo {per_sel}*")
                
                with c_head_der:
                    with st.expander("⚙️ Configurar % de Ponderación para esta Materia", expanded=False):
                        st.caption("Ajusta el peso relativo de cada dimensión. La suma debe dar 100%:")
                        
                        key_p0 = f"peso_0_{mat_sel}_{per_sel}"
                        key_p1 = f"peso_1_{mat_sel}_{per_sel}"
                        key_p2 = f"peso_2_{mat_sel}_{per_sel}"
                        key_p3 = f"peso_3_{mat_sel}_{per_sel}"

                        if key_p0 not in st.session_state: st.session_state[key_p0] = 25
                        if key_p1 not in st.session_state: st.session_state[key_p1] = 25
                        if key_p2 not in st.session_state: st.session_state[key_p2] = 25
                        if key_p3 not in st.session_state: st.session_state[key_p3] = 25

                        p0 = st.slider("🗣️ Participación / Preguntas (%):", 0, 100, st.session_state[key_p0], step=5, key=f"s0_{mat_sel}_{per_sel}")
                        p1 = st.slider("📚 Tareas y Guías (%):", 0, 100, st.session_state[key_p1], step=5, key=f"s1_{mat_sel}_{per_sel}")
                        p2 = st.slider("📝 Evaluaciones Escritas (%):", 0, 100, st.session_state[key_p2], step=5, key=f"s2_{mat_sel}_{per_sel}")
                        p3 = st.slider("🤝 Comportamiento (%):", 0, 100, st.session_state[key_p3], step=5, key=f"s3_{mat_sel}_{per_sel}")

                        st.session_state[key_p0], st.session_state[key_p1], st.session_state[key_p2], st.session_state[key_p3] = p0, p1, p2, p3
                        
                        suma_pesos = p0 + p1 + p2 + p3
                        if suma_pesos == 100:
                            st.success(f"✅ Total Peso: {suma_pesos}% (Configuración Válida)")
                        else:
                            st.warning(f"⚠️ Total Peso: {suma_pesos}% (Ajusta los valores para sumar 100%)")

                pesos_dimensiones = [
                    st.session_state[key_p0],
                    st.session_state[key_p1],
                    st.session_state[key_p2],
                    st.session_state[key_p3]
                ]

                try:
                    res_notas_diarias = supabase.table("notas_diarias").select("*").eq("estudiante_id", est_sel_obj["id"]).eq("materia", mat_sel).eq("periodo", per_sel).execute()
                    notas_guardadas = res_notas_diarias.data or []
                except Exception:
                    notas_guardadas = []

                promedios_dimensiones = []

                for dim_idx, dimension in enumerate(DIMENSIONES_EVALUACION):
                    peso_dim = pesos_dimensiones[dim_idx]
                    
                    with st.expander(f"{dimension}  —  [ Peso: {peso_dim}% ]", expanded=True):
                        st.caption("Ingresa la **Nota**, el **Nombre de la Actividad** y la **Fecha** para cada casilla (C1 - C10):")
                        
                        notas_dim = [n for n in notas_guardadas if n.get("dimension") == dimension]
                        dict_casillas = {n.get("casilla_num"): n for n in notas_dim}
                        
                        val_validos = []

                        for c_num in range(1, 11):
                            n_obj = dict_casillas.get(c_num, {})
                            val_actual = float(n_obj.get("nota", 0.0))
                            act_actual = str(n_obj.get("nombre_actividad", ""))
                            fecha_actual_str = n_obj.get("fecha_actividad")
                            
                            try:
                                f_val = datetime.date.fromisoformat(fecha_actual_str) if fecha_actual_str else datetime.date.today()
                            except Exception:
                                f_val = datetime.date.today()

                            c_n1, c_n2, c_n3 = st.columns([1.2, 2.5, 1.5])
                            
                            with c_n1:
                                val_input = st.number_input(
                                    f"C{c_num} Nota:",
                                    min_value=0.0,
                                    max_value=5.0,
                                    value=val_actual,
                                    step=0.1,
                                    key=f"nota_{est_sel_obj['id']}_{mat_sel}_{per_sel}_{dim_idx}_{c_num}"
                                )
                            with c_n2:
                                act_input = st.text_input(
                                    f"Actividad C{c_num}:",
                                    value=act_actual,
                                    placeholder="Ej. Taller en clase",
                                    key=f"act_{est_sel_obj['id']}_{mat_sel}_{per_sel}_{dim_idx}_{c_num}"
                                )
                            with c_n3:
                                f_input = st.date_input(
                                    f"Fecha C{c_num}:",
                                    value=f_val,
                                    key=f"fec_{est_sel_obj['id']}_{mat_sel}_{per_sel}_{dim_idx}_{c_num}"
                                )

                            f_str_input = f_input.isoformat()
                            if abs(val_input - val_actual) > 0.01 or act_input != act_actual or f_str_input != fecha_actual_str:
                                try:
                                    supabase.table("notas_diarias").upsert({
                                        "estudiante_id": est_sel_obj["id"],
                                        "materia": mat_sel,
                                        "periodo": per_sel,
                                        "dimension": dimension,
                                        "casilla_num": c_num,
                                        "nota": val_input,
                                        "nombre_actividad": act_input.strip(),
                                        "fecha_actividad": f_str_input,
                                        "profesor_id": prof["id"]
                                    }, on_conflict="estudiante_id,materia,periodo,dimension,casilla_num").execute()
                                    st.toast(f"✅ C{c_num} actualizada para {est_sel_nombre}")
                                except Exception as err:
                                    st.error(f"Error al autoguardar C{c_num}: {err}")

                            if val_input > 0.0:
                                val_validos.append(val_input)

                        prom_dim = round(sum(val_validos) / len(val_validos), 2) if val_validos else 0.0
                        promedios_dimensiones.append((prom_dim, peso_dim))
                        
                        col_p1, col_p2 = st.columns([3, 1])
                        with col_p1:
                            st.write(f"**Actividades Calificadas:** {len(val_validos)}/10")
                        with col_p2:
                            st.markdown(f"**Promedio Parcial:** `{prom_dim if prom_dim > 0 else 'S/N'}` *(Aporta {round(prom_dim * (peso_dim / 100), 2)} a la nota final)*")

                st.markdown("---")
                
                acumulado_ponderado = sum(p * (peso / 100) for p, peso in promedios_dimensiones if p > 0.0)
                definitiva_materia = round(acumulado_ponderado, 2) if acumulado_ponderado > 0.0 else 0.0

                if definitiva_materia > 0.0:
                    try:
                        supabase.table("calificaciones").upsert({
                            "estudiante_id": est_sel_obj["id"],
                            "materia": mat_sel,
                            "nota": definitiva_materia,
                            "periodo": per_sel,
                            "profesor_id": prof["id"]
                        }, on_conflict="estudiante_id,materia,periodo").execute()
                    except Exception:
                        pass

                st.markdown(f"### 🏁 Nota Definitiva Calculada ({mat_sel}): **`{definitiva_materia if definitiva_materia > 0 else 'Sin Notas'}`**")
                st.caption("✨ *Sincronizado automáticamente con la pestaña de Boletines Académicos.*")

        # 5. MÓDULO DE BOLETINES ACADÉMICOS
        with t_boletines:
            st.subheader("📄 Generación y Consulta de Boletines Académicos de Escuela Nueva")
            estudiantes = obtener_estudiantes(grado_sel_id, prof["id"])
            
            if not estudiantes:
                st.info("Agrega estudiantes para consultar sus boletines.")
            else:
                dict_e_bol = {e["nombre"]: e for e in estudiantes}
                col_b1, col_b2, col_b3 = st.columns([1.5, 1, 1])
                
                with col_b1:
                    est_bol_nombre = st.selectbox(f"Selecciona Estudiante ({grado_sel_nombre}):", list(dict_e_bol.keys()))
                    est_bol_obj = dict_e_bol[est_bol_nombre]
                with col_b2:
                    per_bol_sel = st.selectbox("Selecciona Periodo:", ["Primero", "Segundo", "Tercero", "Cuarto"])
                with col_b3:
                    sede_bol_txt = st.text_input("Sede Educativa:", value=SEDE_DEFECTO)

                st.markdown("---")

                res_calif = supabase.table("calificaciones").select("*").eq("estudiante_id", est_bol_obj["id"]).eq("periodo", per_bol_sel).execute()
                data_calif = res_calif.data or []
                dict_notas_est = {c["materia"]: c["nota"] for c in data_calif}

                try:
                    res_diarias_todas = supabase.table("notas_diarias").select("*").eq("estudiante_id", est_bol_obj["id"]).eq("periodo", per_bol_sel).execute()
                    diarias_todas = res_diarias_todas.data or []
                except Exception:
                    diarias_todas = []

                if diarias_todas:
                    for m in MATERIAS_LISTA:
                        key_n = normalizar_texto(m)
                        ya_existe = any(normalizar_texto(k) == key_n for k in dict_notas_est.keys())
                        if not ya_existe:
                            notas_m = [nd for nd in diarias_todas if normalizar_texto(nd.get("materia")) == key_n]
                            if notas_m:
                                dims_map = {}
                                for nd in notas_m:
                                    d_nom = nd.get("dimension")
                                    dims_map.setdefault(d_nom, []).append(float(nd.get("nota", 0)))
                                prom_dims = [sum(vals)/len(vals) for vals in dims_map.values() if vals]
                                if prom_dims:
                                    dict_notas_est[m] = round(sum(prom_dims) / len(prom_dims), 2)

                st.write(f"### 📋 Vista Previa de Notas — **{est_bol_nombre}** *(Periodo {per_bol_sel})*")
                
                tabla_boletin = []
                for m in MATERIAS_LISTA:
                    key_norm = normalizar_texto(m)
                    n_val = 0.0
                    for k_materia, val_n in dict_notas_est.items():
                        if normalizar_texto(k_materia) == key_norm:
                            n_val = float(val_n)
                            break

                    val_p, val_n = obtener_valoracion_cualitativa(n_val)
                    tabla_boletin.append({
                        "Área / Asignatura": m,
                        "P.E.N. Valoración Cualitativa": val_p,
                        "Valoración Cuantitativa": f"{n_val:.1f}" if n_val > 0 else "---",
                        "Escala Nacional": val_n
                    })

                df_bol = pd.DataFrame(tabla_boletin)
                st.dataframe(df_bol, use_container_width=True)

                obs_boletin = st.text_area("Observaciones del Docente para este Boletín:", value="Excelente desempeño académico y gran compromiso en las actividades escolares.")

                st.markdown("---")

                pdf_boletin_bytes = generar_pdf_boletin_oficial(
                    estudiante_nombre=est_bol_nombre,
                    grado_nombre=grado_sel_nombre,
                    periodo=per_bol_sel,
                    sede_nombre=sede_bol_txt,
                    profesor_nombre=prof['nombre'],
                    dict_notas=dict_notas_est,
                    observaciones_txt=obs_boletin
                )

                st.download_button(
                    label="📄 Descargar Boletín Académico Oficial en PDF (Formato Oficial 2 Páginas)",
                    data=pdf_boletin_bytes,
                    file_name=f"Boletin_Academico_{est_bol_nombre.replace(' ', '_')}_Periodo_{per_bol_sel}_2026.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary"
                )

        # 6. ADAPTADOR DE GUÍAS DE APRENDIZAJE
        with t_adaptador:
            st.subheader("🧩 Adaptador Automático de Guías a Escuela Nueva")
            st.caption("Carga tus guías tradicionales (PDF, Word o Texto/HTML) y transfórmalas a la secuencia didáctica de 4 Momentos.")

            col_ad1, col_ad2 = st.columns(2)

            with col_ad1:
                st.markdown("#### 1️⃣ Ejes Temáticos y Parámetros")
                mat_guia = st.selectbox("Asignatura de la Guía:", MATERIAS_LISTA, key="mat_guia_sel")
                
                archivo_ejes = st.file_uploader("Subir Documento de Ejes Temáticos (PDF/Word/TXT):", type=["pdf", "docx", "txt"], key="up_ejes")
                
                if archivo_ejes:
                    tipo_e = archivo_ejes.name.split(".")[-1].lower()
                    eje_txt_extraido = extraer_texto_archivo(archivo_ejes, tipo_e)
                    st.success("Eje temático cargado desde el archivo.")
                else:
                    eje_txt_extraido = ""

                eje_tematico_final = st.text_area("O especifica directamente el Eje Temático:", value=eje_txt_extraido if eje_txt_extraido else "La Adición, Términos y Propiedades de la Suma", height=100)

            with col_ad2:
                st.markdown("#### 2️⃣ Guía Original a Transformar")
                origen_tipo = st.radio("Formato de origen:", ["Archivo PDF (.pdf)", "Documento Word (.docx)", "Texto plano / HTML"], horizontal=True)

                texto_guia_original = ""
                
                if origen_tipo == "Archivo PDF (.pdf)":
                    f_pdf = st.file_uploader("Subir Guía en PDF:", type=["pdf"], key="up_pdf_g")
                    if f_pdf:
                        texto_guia_original = extraer_texto_archivo(f_pdf, "pdf")
                elif origen_tipo == "Documento Word (.docx)":
                    f_doc = st.file_uploader("Subir Guía en Word:", type=["docx"], key="up_doc_g")
                    if f_doc:
                        texto_guia_original = extraer_texto_archivo(f_doc, "docx")
                else:
                    texto_guia_original = st.text_area("Pega aquí el texto plano o código HTML de la guía:", height=150)

            st.markdown("---")
            st.markdown("#### 3️⃣ Formato de Salida y Generación")
            
            formato_salida = st.selectbox("Selecciona el formato en el que deseas exportar la guía adaptada:", ["Documento PDF (.pdf)", "Documento Word (.docx)", "Código HTML / Página Web"])

            if st.button("🚀 Transformar y Adaptar Guía a Escuela Nueva", type="primary", use_container_width=True):
                if not texto_guia_original.strip():
                    st.warning("Debes proporcionar el contenido o archivo de la guía original para transformar.")
                else:
                    with st.spinner("Transformando y reestructurando la guía en los 4 Momentos de Escuela Nueva..."):
                        momentos_res = adaptar_contenido_escuela_nueva(texto_guia_original, eje_tematico_final)
                        
                        st.success("¡Guía reestructurada exitosamente!")

                        with st.expander("👁️ Vista Previa de la Guía Adaptada", expanded=True):
                            st.markdown(f"### **GUÍA DE APRENDIZAJE: {mat_guia.upper()}**")
                            st.caption(f"Eje Temático: {eje_tematico_final}")
                            
                            st.markdown("---")
                            st.markdown("#### **🟢 MOMENTO A: Vivencia / Saberes Previos**")
                            st.text(momentos_res["A"])

                            st.markdown("#### **🔵 MOMENTO B: Fundamentación Teórica**")
                            st.text(momentos_res["B"])

                            st.markdown("#### **🟡 MOMENTO C: Ejercitación y Aplicación**")
                            st.text(momentos_res["C"])

                            st.markdown("#### **🟠 MOMENTO D: Evaluación y Extensión**")
                            st.text(momentos_res["D"])

                        st.markdown("---")

                        if formato_salida == "Documento PDF (.pdf)":
                            pdf_g_bytes = generar_pdf_guia_adaptada(mat_guia, eje_tematico_final, momentos_res, prof["nombre"], grado_sel_nombre)
                            st.download_button(
                                label="📄 Descargar Guía Adaptada en PDF",
                                data=pdf_g_bytes,
                                file_name=f"Guia_Escuela_Nueva_{mat_guia.replace(' ', '_')}_{grado_sel_nombre}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                        elif formato_salida == "Documento Word (.docx)":
                            docx_g_bytes = generar_docx_guia_adaptada(mat_guia, eje_tematico_final, momentos_res, prof["nombre"], grado_sel_nombre)
                            st.download_button(
                                label="📝 Descargar Guía Adaptada en Word (.docx)",
                                data=docx_g_bytes,
                                file_name=f"Guia_Escuela_Nueva_{mat_guia.replace(' ', '_')}_{grado_sel_nombre}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        else:
                            html_str = generar_html_guia_adaptada(mat_guia, eje_tematico_final, momentos_res, prof["nombre"], grado_sel_nombre)
                            st.download_button(
                                label="🌐 Descargar Guía Adaptada en HTML",
                                data=html_str.encode("utf-8"),
                                file_name=f"Guia_Escuela_Nueva_{mat_guia.replace(' ', '_')}_{grado_sel_nombre}.html",
                                mime="text/html",
                                use_container_width=True
                            )
                            st.code(html_str, language="html")

        # 7. OBSERVADOR DE CONVIVENCIA CON FILTROS AVANZADOS
        with t_convivencia:
            st.subheader("⚖️ Observador del Estudiante / Convivencia")
            estudiantes = obtener_estudiantes(grado_sel_id, prof["id"])
            
            if estudiantes:
                dict_e = {e["nombre"]: e["id"] for e in estudiantes}
                c_c1, c_c2 = st.columns(2)
                with c_c1:
                    est_conv_nom = st.selectbox("Estudiante a registrar:", list(dict_e.keys()))
                with c_c2:
                    tipo_conv = st.selectbox("Tipo de anotación:", ["Positivo / Reconocimiento", "Llamado de atención", "Falta grave", "Excusa / Justificación"])
                
                desc_conv = st.text_area("Descripción de la situación:")
                if st.button("📝 Guardar Anotación", use_container_width=True):
                    if desc_conv.strip():
                        supabase.table("convivencia").insert({
                            "estudiante_id": dict_e[est_conv_nom],
                            "fecha": datetime.date.today().isoformat(),
                            "tipo": tipo_conv,
                            "descripcion": desc_conv.strip(),
                            "profesor_id": prof["id"]
                        }).execute()
                        st.success("Anotación guardada en el observador.")
                        st.rerun()

                st.markdown("---")
                st.write("### 🔍 Consulta de Observaciones y Excusas (Filtros Rápido)")
                
                res_conv = supabase.table("convivencia").select("*").eq("profesor_id", prof["id"]).execute()
                if res_conv.data:
                    id_to_name = {e["id"]: e["nombre"] for e in estudiantes}
                    
                    df_c = pd.DataFrame([
                        {
                            "Fecha": c["fecha"], 
                            "Estudiante": id_to_name.get(c["estudiante_id"], "Otro Curso"), 
                            "Tipo": c["tipo"], 
                            "Anotación": c["descripcion"]
                        }
                        for c in res_conv.data if c["estudiante_id"] in id_to_name
                    ])
                    
                    col_f1, col_f2 = st.columns(2)
                    with col_f1:
                        filtro_est = st.selectbox("Filtrar por Estudiante:", ["Todos los estudiantes"] + list(dict_e.keys()))
                    with col_f2:
                        filtro_tipo = st.selectbox("Filtrar por Tipo de Anotación:", ["Todos los tipos", "Positivo / Reconocimiento", "Llamado de atención", "Falta grave", "Excusa / Justificación"])

                    if filtro_est != "Todos los estudiantes":
                        df_c = df_c[df_c["Estudiante"] == filtro_est]
                    if filtro_tipo != "Todos los tipos":
                        df_c = df_c[df_c["Tipo"] == filtro_tipo]

                    st.dataframe(df_c, use_container_width=True)
                else:
                    st.info("Aún no hay anotaciones en el observador.")

        # 8. TABLA DE LÍDERES
        with t_lideres:
            st.subheader("🏆 Clasificación General por Puntos y Rachas")
            estudiantes = obtener_estudiantes(grado_sel_id, prof["id"])
            if estudiantes:
                df = pd.DataFrame([
                    {"Estudiante": e["nombre"], "Puntos Acumulados": e["puntos"], "Racha Actual": f"🔥 {e['racha']} días", "Asistencias Totales": e.get("total_asistencias", 0)}
                    for e in estudiantes
                ]).sort_values(by="Puntos Acumulados", ascending=False)
                
                st.dataframe(df, use_container_width=True)
